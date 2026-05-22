from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.main import app

ROOT = Path(__file__).resolve().parents[1]
TMP = ROOT / "_pytest_smoke"
TMP.mkdir(exist_ok=True)


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Name: PyTest Smoke User")
    doc.add_paragraph("Summary: Data analyst with Python, SQL, and dashboards.")
    doc.add_paragraph("Skills: Python, SQL, Pandas, NumPy, Tableau")
    doc.save(path)


def test_analyzer_returns_result_page():
    resume_path = TMP / "pytest_resume.docx"
    _make_docx(resume_path)

    client = TestClient(app)
    client.post("/register", data={"username": "pytest_user", "password": "pytest_pass"})
    client.post("/login", data={"username": "pytest_user", "password": "pytest_pass"})

    with resume_path.open("rb") as handle:
        response = client.post(
            "/analyze?dev_skip_auth=1",
            data={"role": "Data Analyst"},
            files={"resume_file": ("pytest_resume.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    assert response.status_code == 200
    assert b"Analysis complete" in response.content
