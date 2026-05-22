import requests
import docx

# create sample resume docx
doc = docx.Document()
doc.add_paragraph('Name: Test User')
doc.add_paragraph('Summary: Aspiring data analyst with hands-on projects.')
doc.add_paragraph('Skills: Python, SQL, Pandas, NumPy, Tableau')
doc.add_paragraph('Experience: 1 year internship at DataCorp')
doc.save('sample_resume.docx')

s = requests.Session()
BASE = 'http://127.0.0.1:8000'

# Register user (ignore if already exists)
try:
    r = s.post(f'{BASE}/register', data={'username':'testuser','password':'testpass'})
    print('Register status:', r.status_code)
except Exception as e:
    print('Register failed:', e)

# Login
r = s.post(f'{BASE}/login', data={'username':'testuser','password':'testpass'})
print('Login status:', r.status_code)
print('Cookies after login:', s.cookies.get_dict())
print('Login response headers:', dict(r.headers))

# Upload resume and analyze
with open('sample_resume.docx', 'rb') as f:
    files = {'resume_file': ('sample_resume.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {'role': 'Data Analyst'}
    r = s.post(f'{BASE}/analyze?dev_skip_auth=1', files=files, data=data)
    print('Analyze status:', r.status_code)
    print('Analyze response cookies:', s.cookies.get_dict())
    print('Analyze response length:', len(r.content))
    print('Analyze response headers:', dict(r.headers))
    with open('analyze_result.html', 'wb') as out:
        out.write(r.content)
    print('Saved analyze_result.html')
