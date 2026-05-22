from pathlib import Path
import sys
import requests
from docx import Document

BASE = "http://127.0.0.1:8000"
ROOT = Path(__file__).resolve().parents[1]
TMP = ROOT / "_tmp_smoke"
TMP.mkdir(exist_ok=True)


def make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Name: Smoke Test User")
    doc.add_paragraph("Summary: Data analyst with Python, SQL, and dashboards.")
    doc.add_paragraph("Skills: Python, SQL, Pandas, NumPy, Tableau")
    doc.add_paragraph("Experience: Internship at DataCorp")
    doc.save(path)


def main() -> int:
    session = requests.Session()
    resume_path = TMP / "smoke_resume.docx"
    output_path = TMP / "analysis_result.html"
    make_docx(resume_path)

    register = session.post(
        f"{BASE}/register",
        data={"username": "smoke_user", "password": "smoke_pass"},
        timeout=30,
    )
    print(f"register={register.status_code}")

    login = session.post(
        f"{BASE}/login",
        data={"username": "smoke_user", "password": "smoke_pass"},
        timeout=30,
    )
    print(f"login={login.status_code}")

    with resume_path.open("rb") as handle:
        response = session.post(
            f"{BASE}/analyze?dev_skip_auth=1",
            data={"role": "Data Analyst"},
            files={
                "resume_file": (
                    "smoke_resume.docx",
                    handle,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            timeout=30,
        )
    print(f"analyze={response.status_code}")
    output_path.write_bytes(response.content)
    print(f"saved={output_path}")

    if response.status_code != 200 or b"Analysis complete" not in response.content:
        print("smoke test failed", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
