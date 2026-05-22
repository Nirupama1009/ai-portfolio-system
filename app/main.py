"""AI Portfolio Generator - Main Application

A FastAPI-based web application that analyzes resumes against target roles,
computes skill gaps, and generates personalized learning roadmaps.
"""

import logging
import os
import re
from typing import Optional, Dict, Any, List

from fastapi import FastAPI, Request, Form, UploadFile, File, Depends, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session

from .database import SessionLocal, engine, Base
from .models import User
from .services.resume_analyzer import ResumeAnalyzer
from .services.embedding_service import EmbeddingService
from .services.similarity_engine import SimilarityEngine
from .services.roadmap_service import RoadmapService
from .services.resume_extractor import ResumeExtractor
import uuid
try:
    import PyPDF2  # optional dependency for PDF text extraction
except Exception:
    PyPDF2 = None
try:
    import docx as docx_lib  # python-docx for .docx extraction
except Exception:
    docx_lib = None

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# -----------------------
# APP SETUP
# -----------------------
app = FastAPI()
app.add_middleware(SessionMiddleware, secret_key="supersecretkey")

templates = Jinja2Templates(directory="app/templates")

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

Base.metadata.create_all(bind=engine)

embedding_service = EmbeddingService()
similarity_engine = SimilarityEngine()

# -----------------------
# ROLE DESCRIPTIONS (🔥 FIX ADDED)
# -----------------------
ROLE_DESCRIPTIONS = {
    "Data Analyst": "Python SQL Pandas NumPy Power BI Tableau Excel Statistics",
    "Machine Learning Engineer": "Python Machine Learning Deep Learning TensorFlow Scikit-learn",
    "Data Scientist": "Python Statistics Machine Learning Pandas NumPy",
    "Backend Developer": "Python FastAPI Django SQL APIs",
    "AI Engineer": "Deep Learning NLP Computer Vision TensorFlow",
    "Business Intelligence Analyst": "SQL Power BI Tableau Reporting",
    "Full Stack Developer": "HTML CSS JavaScript Python Django",
    "Data Engineer": "Python SQL ETL Data Warehousing",
    "NLP Engineer": "Transformers BERT NLP Python",
    "Cloud Engineer": "AWS Azure GCP Docker Kubernetes"
}

PORTFOLIO_TEMPLATES = [
    {
        "slug": "data-analyst",
        "title": "Data Analyst",
        "tagline": "Clean dashboards and measurable impact",
        "accent": "#5cc8ff",
        "summary": "A data-focused layout with sharp metrics, charts, and project evidence.",
    },
    {
        "slug": "backend-developer",
        "title": "Backend Developer",
        "tagline": "Structured, technical, and API-first",
        "accent": "#7c9cff",
        "summary": "Best for engineers who want to highlight systems, reliability, and architecture.",
    },
    {
        "slug": "full-stack-developer",
        "title": "Full Stack Developer",
        "tagline": "Modern product work with balanced depth",
        "accent": "#3dd6c5",
        "summary": "Ideal for product builders with front-end polish and backend strength.",
    },
    {
        "slug": "machine-learning-engineer",
        "title": "Machine Learning Engineer",
        "tagline": "Research-to-production portfolio",
        "accent": "#ffb84d",
        "summary": "Showcase models, experiments, metrics, and deployment work.",
    },
    {
        "slug": "business-intelligence-analyst",
        "title": "BI Analyst",
        "tagline": "Reporting, decisions, and business clarity",
        "accent": "#f26d78",
        "summary": "A business-ready layout for KPIs, insights, and stakeholder impact.",
    },
    {
        "slug": "data-engineer",
        "title": "Data Engineer",
        "tagline": "Pipelines, scale, and dependable systems",
        "accent": "#9d7cff",
        "summary": "Focused on ETL, orchestration, and data platform outcomes.",
    },
]

PORTFOLIO_TEMPLATE_LAYOUTS = {
    "data-analyst": "analytics",
    "business-intelligence-analyst": "analytics",
    "machine-learning-engineer": "research",
    "ai-engineer": "research",
    "nlp-engineer": "research",
    "backend-developer": "builder",
    "full-stack-developer": "builder",
    "data-engineer": "builder",
}

# -----------------------
# DATABASE DEPENDENCY
# -----------------------
def get_db() -> Any:
    """Dependency to get database session."""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# -----------------------
# LOGOUT
# -----------------------
@app.get("/logout")
def logout(request: Request) -> RedirectResponse:
    """Handle user logout."""
    username = request.session.get("username", "Unknown")
    request.session.clear()
    logger.info(f"User {username} logged out")
    return RedirectResponse("/login", status_code=302)

# -----------------------
# ROOT
# -----------------------
@app.get("/", response_class=HTMLResponse)
def root(request: Request) -> RedirectResponse:
    """Root endpoint - redirects to login or analyzer based on session."""
    if "user_id" not in request.session:
        logger.info("Redirecting to login")
        return RedirectResponse("/login", status_code=302)
    logger.info(f"User {request.session.get('username')} accessing analyzer")
    return RedirectResponse("/analyzer", status_code=302)

# -----------------------
# LOGIN PAGE
# -----------------------
@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request) -> HTMLResponse:
    """Display login page."""
    return templates.TemplateResponse(request, "login.html", {"request": request})

@app.post("/login")
def login(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    """Handle login authentication with validation."""
    if not username or not username.strip():
        logger.warning("Login attempt with empty username")
        return templates.TemplateResponse(
            request, "login.html",
            {"request": request, "error": "Username cannot be empty"}
        )
    
    user = db.query(User).filter(User.username == username).first()
    if user and user.password == password:
        request.session["user_id"] = user.id
        request.session["username"] = user.username
        logger.info(f"User {username} logged in successfully")
        return RedirectResponse("/", status_code=302)
    
    logger.warning(f"Failed login attempt for user: {username}")
    return templates.TemplateResponse(request, "login.html", {"request": request, "error": "Invalid credentials"})

# -----------------------
# REGISTER PAGE
# -----------------------
@app.get("/register", response_class=HTMLResponse)
def register_page(request: Request) -> HTMLResponse:
    """Display registration page."""
    return templates.TemplateResponse(request, "register.html", {"request": request})

@app.post("/register")
def register(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    """Handle user registration with validation."""
    if not username or len(username) < 3:
        logger.warning("Registration attempt with invalid username")
        return templates.TemplateResponse(
            request, "register.html",
            {"request": request, "error": "Username must be at least 3 characters"}
        )
    
    if not password or len(password) < 6:
        logger.warning(f"Registration attempt with weak password for user: {username}")
        return templates.TemplateResponse(
            request, "register.html",
            {"request": request, "error": "Password must be at least 6 characters"}
        )
    
    existing_user = db.query(User).filter(User.username == username).first()
    if existing_user:
        logger.warning(f"Registration attempt for existing user: {username}")
        return templates.TemplateResponse(
            request, "register.html",
            {"request": request, "error": "Username already exists"}
        )
    
    try:
        user = User(username=username, password=password)
        db.add(user)
        db.commit()
        logger.info(f"New user registered: {username}")
        return RedirectResponse("/login", status_code=302)
    except Exception as e:
        db.rollback()
        logger.error(f"Registration error for user {username}: {str(e)}")
        return templates.TemplateResponse(
            request, "register.html",
            {"request": request, "error": "Registration failed. Please try again."}
        )

# -----------------------
# ANALYZER PAGE
# -----------------------
@app.get("/analyzer", response_class=HTMLResponse)
def analyzer_page(request: Request):
    """Display resume analyzer page."""
    if "user_id" not in request.session:
        logger.warning("Unauthorized access to analyzer")
        return RedirectResponse("/login", status_code=302)
    return templates.TemplateResponse(request, "analyze.html", {"request": request})

# -----------------------
# EXTRACT RESUME DETAILS
# -----------------------
@app.post("/extract-resume")
async def extract_resume(resume_file: UploadFile = File(...)):
    """Extract only the basic identity details from a resume file."""
    try:
        if not resume_file:
            return {"error": "No file provided"}
        
        file_ext = os.path.splitext(resume_file.filename)[1].lower()
        file_content = await resume_file.read()
        
        # Extract text from file
        resume_text = ""
        if file_ext == ".txt":
            resume_text = file_content.decode('utf-8', errors='ignore')
        elif file_ext == ".pdf" and PyPDF2 is not None:
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                resume_text += page.extract_text() + "\n"
        elif file_ext == ".docx" and docx_lib is not None:
            import io
            doc = docx_lib.Document(io.BytesIO(file_content))
            resume_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_ext in [".doc", ".docx"] and docx_lib is not None:
            import io
            try:
                doc = docx_lib.Document(io.BytesIO(file_content))
                resume_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception:
                resume_text = file_content.decode('utf-8', errors='ignore')
        else:
            resume_text = file_content.decode('utf-8', errors='ignore')
        
        # Extract structured data
        extractor = ResumeExtractor(resume_text)
        extracted = extractor.extract_all()
        
        return {
            "success": True,
            "name": extracted.get("name"),
            "email": extracted.get("email"),
            "linkedin": extracted.get("linkedin"),
            "github": extracted.get("github"),
            "technical_skills": extracted.get("technical_skills", []),
            "soft_skills": extracted.get("soft_skills", []),
        }
    except Exception as e:
        logger.error(f"Error extracting resume: {str(e)}")
        return {"error": str(e), "success": False}

# -----------------------
# ANALYZE
# -----------------------
@app.post("/analyze", response_class=HTMLResponse)
async def analyze(request: Request, resume_text: str = Form(None), role: str = Form(None), resume_file: UploadFile = File(None)):
    """Analyze resume against target role with skill gap detection."""
    # Authentication check (allow dev bypass with query param ?dev_skip_auth=1)
    if request.query_params.get('dev_skip_auth') != '1':
        if "user_id" not in request.session:
            logger.warning("Unauthorized access to analyze endpoint")
            return RedirectResponse("/login", status_code=302)

    # Require a resume file upload (pasted text removed)
    if not (resume_file and resume_file.filename):
        logger.warning("Analyze request missing resume file")
        return templates.TemplateResponse(request, "analyze.html", {"request": request, "error": "Please upload a resume file (PDF or TXT)."})

    # Read and extract text from uploaded file
    try:
        content = await resume_file.read()
        filename = os.path.basename(resume_file.filename)
        saved_path = os.path.join(UPLOAD_FOLDER, filename)
        with open(saved_path, 'wb') as f:
            f.write(content)
        logger.info(f"Saved uploaded resume to {saved_path}")

        file_ext = os.path.splitext(filename)[1].lower()
        extracted_text = ''
        if file_ext == '.pdf' and PyPDF2 is not None:
            try:
                reader = PyPDF2.PdfReader(saved_path)
                for page in reader.pages:
                    text = page.extract_text() or ''
                    extracted_text += text + '\n'
            except Exception:
                extracted_text = content.decode('utf-8', errors='ignore')
        elif file_ext in ('.txt',):
            extracted_text = content.decode('utf-8', errors='ignore')
        elif file_ext == '.docx' and docx_lib is not None:
            try:
                doc = docx_lib.Document(saved_path)
                for para in doc.paragraphs:
                    extracted_text += (para.text or '') + '\n'
            except Exception:
                extracted_text = content.decode('utf-8', errors='ignore')
        else:
            # fallback for other formats (docx decoding optional)
            try:
                extracted_text = content.decode('utf-8', errors='ignore')
            except Exception:
                extracted_text = ''

        resume_text = extracted_text or ''
    except Exception as e:
        logger.error(f"Error reading uploaded resume: {e}")
        return templates.TemplateResponse(request, "analyze.html", {"request": request, "error": "Unable to read uploaded file."})

    if not role or role not in ROLE_DESCRIPTIONS:
        logger.warning(f"Analyze request with invalid role: {role}")
        return templates.TemplateResponse(request, "analyze.html", {"request": request, "error": "Please select a valid role"})

    try:
        current_user = {"username": request.session.get("username")}
        
        # Extract skills
        analyzer = ResumeAnalyzer()
        skills = analyzer.extract_skills(resume_text)
        
        # Get role skills
        role_text = ROLE_DESCRIPTIONS.get(role, "")
        
        # Calculate similarity
        resume_vector = embedding_service.get_embedding(resume_text)
        role_vector = embedding_service.get_embedding(role_text)
        
        similarity_score = similarity_engine.calculate_similarity(
            resume_vector,
            role_vector
        )
        
        # Calculate metrics
        match_percentage = round(similarity_score * 100, 2)
        skill_gap = round(100 - match_percentage, 2)
        
        # Find missing skills (tokenize role description and compare)
        role_tokens = re.findall(r"[a-z0-9]+", role_text.lower())
        detected_skills = {s.lower() for s in skills}

        missing_skills = [
            token.capitalize() for token in role_tokens if token not in detected_skills
        ]
        
        # Generate learning roadmap
        roadmap_service = RoadmapService()
        learning_roadmap = roadmap_service.generate_roadmap(missing_skills)
        
        logger.info(f"Analysis completed for user {current_user['username']}: {role} - {match_percentage}% match")
        
        return templates.TemplateResponse(
            request,
            "result.html",
            {
                "request": request,
                "current_user": current_user,
                "resume_text": resume_text,
                "role": role,
                "skills": skills,
                "missing_skills": missing_skills,
                "match_percentage": match_percentage,
                "skill_gap": skill_gap,
                "learning_roadmap": learning_roadmap
            }
        )
    except Exception as e:
        logger.error(f"Error during analysis: {str(e)}")
        return templates.TemplateResponse(
            request, "analyze.html",
            {"request": request, "error": "An error occurred during analysis. Please try again."}
        )

# -----------------------
# PORTFOLIO FORM ROUTE
# -----------------------
@app.get("/portfolio-form", response_class=HTMLResponse)
async def portfolio_form(
    request: Request,
    resume_text: Optional[str] = None,
    role: Optional[str] = None,
    template: Optional[str] = None,
):
    """Display portfolio generation form."""
    selected_template = next((item for item in PORTFOLIO_TEMPLATES if item["slug"] == template), None)
    selected_template_layout = PORTFOLIO_TEMPLATE_LAYOUTS.get(template, "classic")
    return templates.TemplateResponse(
        request,
        "portfolio-form.html",
        {
            "request": request,
            "resume_text": resume_text or "",
            "role": role or "",
            "template": template or "",
            "selected_template": selected_template,
            "selected_template_layout": selected_template_layout,
        }
    )


@app.get("/portfolio-templates", response_class=HTMLResponse)
async def portfolio_templates(request: Request, resume_text: Optional[str] = None, role: Optional[str] = None):
    """Show portfolio template options before the detailed form."""
    return templates.TemplateResponse(
        request,
        "portfolio_templates.html",
        {
            "request": request,
            "resume_text": resume_text or "",
            "role": role or "",
            "templates": PORTFOLIO_TEMPLATES,
        }
    )


@app.get("/quick-portfolio", response_class=HTMLResponse)
async def quick_portfolio_form(request: Request):
    """Display quick portfolio form accessible without login."""
    return templates.TemplateResponse(request, "quick_portfolio_form.html", {"request": request})


@app.post("/quick-portfolio", response_class=HTMLResponse)
async def quick_generate_portfolio(request: Request):
    """Generate portfolio page from the quick form."""
    form = await request.form()
    name = form.get("name", "")
    email = form.get("email", "")
    linkedin = form.get("linkedin", "")
    github = form.get("github", "")
    summary = form.get("summary", "")
    technical_skills = form.get("technical_skills", "")
    soft_skills = form.get("soft_skills", "")
    school_education = form.get("school_education", "")
    college_education = form.get("college_education", "")
    experience = form.get("experience", "")
    project_names = form.getlist("project_name[]") if hasattr(form, 'getlist') else []
    project_descs = form.getlist("project_desc[]") if hasattr(form, 'getlist') else []

    # profile image handling
    profile_image = form.get("profile_image")
    profile_image_size = form.get('profile_image_size') or None
    profile_image_shape = form.get('profile_image_shape') or 'circle'
    profile_image_align = form.get('profile_image_align') or 'center'
    image_filename = None
    try:
        if profile_image and getattr(profile_image, 'filename', None):
            allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
            file_ext = os.path.splitext(profile_image.filename)[1].lower()
            if file_ext in allowed_extensions:
                image_filename = os.path.basename(profile_image.filename)
                file_path = os.path.join(UPLOAD_FOLDER, image_filename)
                with open(file_path, 'wb') as f:
                    f.write(await profile_image.read())
    except Exception:
        image_filename = None

    project_list = []
    for i, name_p in enumerate(project_names or []):
        title = name_p.strip()
        desc = (project_descs[i].strip() if i < len(project_descs) else '').strip()
        if title or desc:
            project_list.append({"title": title, "description": desc})

    # Certificates
    certificate_names = form.getlist("certificate_name[]") if hasattr(form, 'getlist') else []
    certificate_descs = form.getlist("certificate_desc[]") if hasattr(form, 'getlist') else []
    certificate_list = []
    for i, cname in enumerate(certificate_names or []):
        ctitle = cname.strip()
        cdesc = (certificate_descs[i].strip() if i < len(certificate_descs) else '').strip()
        if ctitle or cdesc:
            certificate_list.append({"title": ctitle, "issuer": cdesc})

    return templates.TemplateResponse(
        request,
        "portfolio.html",
        {
            "request": request,
            "name": name,
            "email": email,
            "linkedin": linkedin,
            "github": github,
            "summary": summary,
            "technical_skills": technical_skills,
                "soft_skills": soft_skills,
            "school_education": school_education,
            "college_education": college_education,
            "experience": experience,
            "role": "",
            "template_accent": form.get('template_accent') or None,
            "profile_image_size": form.get('profile_image_size') or None,
            "profile_image_shape": form.get('profile_image_shape') or 'circle',
            "profile_image_align": form.get('profile_image_align') or 'center',
            "projects": project_list,
                "certificates": certificate_list,
            "image_path": image_filename
        }
    )

# -----------------------
# PORTFOLIO ROUTE
# -----------------------
@app.post("/portfolio", response_class=HTMLResponse)
async def generate_portfolio(request: Request, profile_image: UploadFile = File(None)):
    """Generate professional portfolio page with file upload."""
    # Input validation
    # For basic design revert, accept posted form fields directly from request.form()
    form = await request.form()
    name = form.get("name", "")
    email = form.get("email", "")
    linkedin = form.get("linkedin", "")
    github = form.get("github", "")
    summary = form.get("summary", "")
    technical_skills = form.get("technical_skills", "")
    soft_skills = form.get("soft_skills", "")
    school_education = form.get("school_education", "")
    college_education = form.get("college_education", "")
    experience = form.get("experience", "")
    role = form.get("role", "")
    portfolio_template = form.get("portfolio_template", "")
    selected_template = next((item for item in PORTFOLIO_TEMPLATES if item["slug"] == portfolio_template), None)
    portfolio_template_label = selected_template["title"] if selected_template else portfolio_template
    portfolio_template_layout = PORTFOLIO_TEMPLATE_LAYOUTS.get(portfolio_template, "classic")
    # projects are sent as arrays `project_name[]` and `project_desc[]`
    project_names = form.getlist("project_name[]") if hasattr(form, 'getlist') else []
    project_descs = form.getlist("project_desc[]") if hasattr(form, 'getlist') else []
    profile_image = form.get("profile_image")
    profile_image_size = form.get("profile_image_size") or None
    profile_image_shape = form.get("profile_image_shape") or 'circle'
    profile_image_align = form.get("profile_image_align") or 'center'

    try:
        image_filename = None
        if profile_image and profile_image.filename:
            allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
            file_ext = os.path.splitext(profile_image.filename)[1].lower()
            if file_ext in allowed_extensions:
                image_filename = os.path.basename(profile_image.filename)
                file_path = os.path.join(UPLOAD_FOLDER, image_filename)
                with open(file_path, 'wb') as f:
                    f.write(await profile_image.read())
                logger.info(f"Saved profile image to {file_path}")
            else:
                logger.warning(f"Profile image has unsupported extension: {file_ext}")

        project_list = []
        for i, name_p in enumerate(project_names or []):
            title = name_p.strip()
            desc = (project_descs[i].strip() if i < len(project_descs) else '').strip()
            if title or desc:
                project_list.append({"title": title, "description": desc})
        # Certificates
        certificate_names = form.getlist("certificate_name[]") if hasattr(form, 'getlist') else []
        certificate_descs = form.getlist("certificate_desc[]") if hasattr(form, 'getlist') else []
        certificate_list = []
        for i, cname in enumerate(certificate_names or []):
            ctitle = cname.strip()
            cdesc = (certificate_descs[i].strip() if i < len(certificate_descs) else '').strip()
            if ctitle or cdesc:
                certificate_list.append({"title": ctitle, "issuer": cdesc})
        template_accent = form.get('template_accent') or (selected_template.get('accent') if selected_template else None)
        return templates.TemplateResponse(
            request,
            "portfolio.html",
            {
                "request": request,
                "name": name,
                "email": email,
                "linkedin": linkedin,
                "github": github,
                "summary": summary,
                "technical_skills": technical_skills,
                    "soft_skills": soft_skills,
                "school_education": school_education,
                "college_education": college_education,
                "experience": experience,
                "role": role,
                "portfolio_template": portfolio_template,
                "portfolio_template_label": portfolio_template_label,
                "portfolio_template_layout": portfolio_template_layout,
                "profile_image_size": profile_image_size,
                "profile_image_shape": profile_image_shape,
                "profile_image_align": profile_image_align,
                "template_accent": template_accent,
                "projects": project_list,
                "certificates": certificate_list,
                "image_path": image_filename
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating portfolio: {str(e)}")
        raise HTTPException(status_code=500, detail="Error generating portfolio. Please try again.")


@app.post("/preview-portfolio", response_class=HTMLResponse)
async def preview_portfolio(request: Request, profile_image: UploadFile = File(None)):
    """Render a server-side preview of the portfolio in a new tab/window."""
    form = await request.form()
    try:
        name = form.get("name", "")
        email = form.get("email", "")
        linkedin = form.get("linkedin", "")
        github = form.get("github", "")
        summary = form.get("summary", "")
        technical_skills = form.get("technical_skills", "")
        soft_skills = form.get("soft_skills", "")
        school_education = form.get("school_education", "")
        college_education = form.get("college_education", "")
        experience = form.get("experience", "")
        role = form.get("role", "")
        portfolio_template = form.get("portfolio_template", "")
        selected_template = next((item for item in PORTFOLIO_TEMPLATES if item["slug"] == portfolio_template), None)
        portfolio_template_label = selected_template["title"] if selected_template else portfolio_template
        portfolio_template_layout = PORTFOLIO_TEMPLATE_LAYOUTS.get(portfolio_template, "classic")

        # save preview image if provided
        image_filename = None
        if profile_image and getattr(profile_image, 'filename', None):
            file_ext = os.path.splitext(profile_image.filename)[1].lower()
            allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
            if file_ext in allowed_extensions:
                # prefix with uuid to avoid clashes
                image_filename = f"preview_{uuid.uuid4().hex}_{os.path.basename(profile_image.filename)}"
                file_path = os.path.join(UPLOAD_FOLDER, image_filename)
                with open(file_path, 'wb') as f:
                    f.write(await profile_image.read())

        # projects and certificates
        project_names = form.getlist("project_name[]") if hasattr(form, 'getlist') else []
        project_descs = form.getlist("project_desc[]") if hasattr(form, 'getlist') else []
        project_list = []
        for i, name_p in enumerate(project_names or []):
            title = name_p.strip()
            desc = (project_descs[i].strip() if i < len(project_descs) else '').strip()
            if title or desc:
                project_list.append({"title": title, "description": desc})

        certificate_names = form.getlist("certificate_name[]") if hasattr(form, 'getlist') else []
        certificate_descs = form.getlist("certificate_desc[]") if hasattr(form, 'getlist') else []
        certificate_list = []
        for i, cname in enumerate(certificate_names or []):
            ctitle = cname.strip()
            cdesc = (certificate_descs[i].strip() if i < len(certificate_descs) else '').strip()
            if ctitle or cdesc:
                certificate_list.append({"title": ctitle, "issuer": cdesc})

        template_accent = form.get('template_accent') or (selected_template.get('accent') if selected_template else None)
        profile_image_size = form.get('profile_image_size') or None
        profile_image_shape = form.get('profile_image_shape') or 'circle'
        profile_image_align = form.get('profile_image_align') or 'center'

        return templates.TemplateResponse(
            request,
            "portfolio.html",
            {
                "request": request,
                "name": name,
                "email": email,
                "linkedin": linkedin,
                "github": github,
                "summary": summary,
                "technical_skills": technical_skills,
                "soft_skills": soft_skills,
                "school_education": school_education,
                "college_education": college_education,
                "experience": experience,
                "role": role,
                "portfolio_template": portfolio_template,
                "portfolio_template_label": portfolio_template_label,
                "portfolio_template_layout": portfolio_template_layout,
                "projects": project_list,
                "certificates": certificate_list,
                "image_path": image_filename,
                "template_accent": template_accent,
                "profile_image_size": profile_image_size,
                "profile_image_shape": profile_image_shape,
                "profile_image_align": profile_image_align,
            }
        )
    except Exception as e:
        logger.error(f"Error rendering preview: {e}")
        return Response(content=str(e), status_code=500)